import streamlit as st
import os
import json
import base64
import io
import pandas as pd
import docx
import PyPDF2
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import openai
from datetime import datetime
import re
import csv

# Set page configuration
st.set_page_config(
    page_title="Meeting Minutes Analyzer", 
    layout="wide",
    page_icon="📝",
    initial_sidebar_state="expanded"
)

# Custom CSS for better UI
st.markdown("""
<style>
    .main .block-container {padding-top: 2rem;}
    h1, h2, h3 {color: #1E3A8A;}
    .stTabs [data-baseweb="tab-list"] {gap: 8px;}
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: #F3F4F6;
        border-radius: 4px 4px 0px 0px;
        gap: 1px;
        padding-left: 16px;
        padding-right: 16px;
    }
    .stTabs [aria-selected="true"] {
        background-color: #DBEAFE !important;
        color: #1E40AF !important;
        font-weight: bold;
    }
    div.stButton > button {
        background-color: #1E40AF;
        color: white;
        font-weight: bold;
        padding: 0.5rem 1rem;
        border-radius: 0.5rem;
    }
    div.stButton > button:hover {
        background-color: #1E3A8A;
        border-color: #1E3A8A;
    }
    .success-box {
        padding: 1rem;
        background-color: #ECFDF5;
        border-left: 5px solid #10B981;
        border-radius: 0.5rem;
        margin-bottom: 1rem;
    }
    .task-box {
        padding: 1rem;
        background-color: #F3F4F6;
        border-radius: 0.5rem;
        margin-bottom: 0.5rem;
    }
    .summary-box {
        padding: 1rem;
        background-color: #EFF6FF;
        border-radius: 0.5rem;
        margin-bottom: 1rem;
        border-left: 5px solid #3B82F6;
    }
    .css-1544g2n.e1fqkh3o4 {
        padding-top: 2rem;
    }
</style>
""", unsafe_allow_html=True)

# Constants
SCOPES = ['https://www.googleapis.com/auth/gmail.readonly', 
          'https://www.googleapis.com/auth/gmail.send']
CREDENTIALS_FILE = "credentials.json"
TOKEN_FILE = "token.json"

# Initialize session state variables
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False
if 'service' not in st.session_state:
    st.session_state.service = None
if 'auth_code' not in st.session_state:
    st.session_state.auth_code = None
if 'auth_url' not in st.session_state:
    st.session_state.auth_url = None
if 'flow' not in st.session_state:
    st.session_state.flow = None
if 'transcript_content' not in st.session_state:
    st.session_state.transcript_content = None
if 'participants' not in st.session_state:
    st.session_state.participants = None
if 'summary' not in st.session_state:
    st.session_state.summary = None
if 'tasks' not in st.session_state:
    st.session_state.tasks = None

# OpenAI API key setup (use a more secure approach in production)
openai_api_key = st.secrets["key"]
client = openai.OpenAI(api_key=openai_api_key)

# Gmail API credentials
gmail_credentials = {
    "web": {
        "client_id": st.secrets["id"],
        "project_id": st.secrets["project_id"],
        "auth_uri": "https://accounts.google.com/o/oauth2/auth",
        "token_uri": "https://oauth2.googleapis.com/token",
        "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
        "client_secret": st.secrets["sec"],
        "redirect_uris":["http://meeting-monitor-k5i6nn5kscxpo7kbhhvjt4.streamlit.app"]
    }
}

# Save credentials to a file
def save_credentials():
    with open(CREDENTIALS_FILE, "w") as f:
        json.dump(gmail_credentials, f)

# Modified authentication function to use the Streamlit UI
def start_auth_flow():
    save_credentials()
    # Create a flow instance with the redirect URI set to localhost:8501
    flow = InstalledAppFlow.from_client_secrets_file(
        CREDENTIALS_FILE, 
        SCOPES, 
        redirect_uri="http://meeting-monitor-k5i6nn5kscxpo7kbhhvjt4.streamlit.app"
    )
    
    # Generate the authorization URL
    auth_url, _ = flow.authorization_url(
        access_type='offline',
        include_granted_scopes='true'
    )
    
    # Store the flow in the session state
    st.session_state.flow = flow
    st.session_state.auth_url = auth_url
    
    return auth_url

# Check for authorization code in URL - FIXED with st.query_params
def check_url_for_auth_code():
    try:
        # Get the query parameters from URL with the updated API
        query_params = st.query_params
        
        # Check if 'code' is in the query parameters
        if 'code' in query_params:
            auth_code = query_params['code']
            st.session_state.auth_code = auth_code
            return auth_code
        return None
    except Exception as e:
        st.error(f"Error accessing query parameters: {str(e)}")
        return None

# Function to complete the auth flow with the code
def complete_auth_flow(code):
    try:
        # Exchange the authorization code for credentials
        if st.session_state.flow is None:
            st.error("Authentication flow not initialized. Please start authentication process again.")
            return None
            
        st.session_state.flow.fetch_token(code=code)
        creds = st.session_state.flow.credentials
        
        # Save the credentials for future use
        with open(TOKEN_FILE, "w") as token:
            token.write(creds.to_json())
        
        return creds
    except Exception as e:
        st.error(f"Error completing authentication: {str(e)}")
        return None

# Check if token exists and is valid
def get_credentials():
    if os.path.exists(TOKEN_FILE):
        try:
            creds = Credentials.from_authorized_user_info(
                json.loads(open(TOKEN_FILE).read())
            )
            if not creds.expired:
                return creds
        except Exception as e:
            st.error(f"Error loading credentials: {str(e)}")
            if os.path.exists(TOKEN_FILE):
                os.remove(TOKEN_FILE)
    return None

# Get Gmail service
def get_gmail_service():
    creds = get_credentials()
    if creds:
        return build('gmail', 'v1', credentials=creds)
    return None

# Send an email
def send_email(service, to, subject, body_html):
    message = MIMEMultipart('alternative')
    message['to'] = to
    message['subject'] = subject
    
    # Create HTML part
    html_part = MIMEText(body_html, 'html')
    message.attach(html_part)
    
    raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode('utf-8')
    
    try:
        sent_message = service.users().messages().send(
            userId='me',
            body={'raw': raw_message}
        ).execute()
        return True, sent_message
    except Exception as e:
        return False, str(e)

# Function to read transcript from uploaded file
def read_transcript(uploaded_file):
    if uploaded_file is None:
        return None
    
    file_extension = uploaded_file.name.split('.')[-1].lower()
    content = ""
    
    try:
        if file_extension == 'txt':
            content = uploaded_file.getvalue().decode('utf-8')
        elif file_extension == 'docx':
            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
            content = "\n".join([para.text for para in doc.paragraphs])
        elif file_extension == 'pdf':
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.getvalue()))
            for page_num in range(len(pdf_reader.pages)):
                content += pdf_reader.pages[page_num].extract_text() + "\n"
        else:
            return None
        
        return content
    except Exception as e:
        st.error(f"Error reading file: {str(e)}")
        return None

# Function to read participants from uploaded file - FIXED to properly handle CSV parsing
def read_participants(uploaded_file):
    if uploaded_file is None:
        return None
    
    file_extension = uploaded_file.name.split('.')[-1].lower()
    participants = []
    
    try:
        if file_extension == 'csv':
            # Read CSV data and skip any initial blank rows
            csv_data = uploaded_file.getvalue().decode('utf-8').splitlines()
            # Use Python's CSV reader for more robust parsing
            reader = csv.DictReader(csv_data)
            
            required_columns = ['name', 'email', 'expertise']
            
            # Check if all required columns are present (case-insensitive)
            header_lower = [col.lower() for col in reader.fieldnames]
            if not all(col.lower() in header_lower for col in required_columns):
                st.error("CSV file must contain 'name', 'email', and 'expertise' columns")
                return None
            
            # Map actual column names to expected column names (handling case differences)
            col_mapping = {}
            for req_col in required_columns:
                for actual_col in reader.fieldnames:
                    if actual_col.lower() == req_col.lower():
                        col_mapping[req_col] = actual_col
            
            # Convert CSV rows to list of dictionaries
            for row in reader:
                if any(row.values()):  # Skip completely empty rows
                    participant = {
                        'name': row[col_mapping['name']].strip(),
                        'email': row[col_mapping['email']].strip(),
                        'expertise': row[col_mapping['expertise']].strip()
                    }
                    participants.append(participant)
                
        elif file_extension == 'txt':
            content = uploaded_file.getvalue().decode('utf-8')
            lines = content.split('\n')
            for line in lines:
                if line.strip():  # Skip empty lines
                    parts = line.split(',')
                    if len(parts) >= 3:
                        participants.append({
                            'name': parts[0].strip(),
                            'email': parts[1].strip(),
                            'expertise': parts[2].strip()
                        })
                    else:
                        st.error(f"Invalid format in line: {line}. Expected 'name, email, expertise'")
                        return None
                        
        elif file_extension == 'docx':
            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty lines
                    parts = para.text.split(',')
                    if len(parts) >= 3:
                        participants.append({
                            'name': parts[0].strip(),
                            'email': parts[1].strip(),
                            'expertise': parts[2].strip()
                        })
                    else:
                        st.error(f"Invalid format in line: {para.text}. Expected 'name, email, expertise'")
                        return None
        else:
            return None
        
        return participants
    except Exception as e:
        st.error(f"Error reading participants file: {str(e)}")
        return None
# Modified functions for improved meeting transcript analysis

def generate_meeting_summary(transcript):
    if not transcript:
        return "No transcript provided for summarization."
    
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a professional assistant that creates concise yet comprehensive summaries of meeting transcripts."},
                {"role": "user", "content": f"Please provide a summarized version of this meeting transcript that captures the key points, decisions, and overall purpose:\n\n{transcript}"}
            ],
            max_tokens=500
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error generating meeting summary: {str(e)}"

# Modified task extraction to avoid creating artificial tasks
def extract_tasks_and_assign(transcript, participants):
    if not transcript or not participants:
        return {"tasks": []}
    
    # Create a set of lowercase participant names for strict matching
    participant_names = {p['name'].lower() for p in participants}
    
    # Create a mapping of name to email and expertise
    participant_info = ""
    name_to_email = {}
    name_to_expertise = {}
    
    for p in participants:
        name_lower = p['name'].lower()
        name_to_email[name_lower] = p['email']
        name_to_expertise[name_lower] = p['expertise']
        participant_info += f"- {p['name']}: {p['expertise']}, Email: {p['email']}\n"
    
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": """You are a professional assistant that identifies ONLY explicitly mentioned tasks from meeting transcripts.
                
                STRICT RULES:
                1. ONLY extract tasks that are EXPLICITLY mentioned in the transcript.
                2. NEVER invent, infer, or create tasks that aren't clearly stated in the transcript.
                3. ONLY assign tasks to people who are EXPLICITLY mentioned as responsible in the transcript AND appear in the provided participant list.
                4. If a task exists but has no clear assignee, mark it as 'Unassigned'.
                5. If no tasks are mentioned at all, return an empty tasks array.
                6. Do not try to be helpful by creating tasks - only report what's in the transcript."""},
                
                {"role": "user", "content": f"""Based on the meeting transcript below, identify ONLY explicitly mentioned tasks and action items.
                
                IMPORTANT CONSTRAINTS:
                - Task extraction should be CONSERVATIVE - only include tasks with clear action verbs and deliverables.
                - A person can only be assigned a task if they are EXPLICITLY mentioned as responsible AND they appear in the team member list below.
                - Return a COMPLETELY EMPTY tasks array if no explicit tasks are mentioned.
                
                Format your response as a JSON object with a 'tasks' array. Each task should include:
                - 'task': The specific action item mentioned (verbatim from transcript when possible)
                - 'assignee': The person explicitly assigned (must match a name in team list) or 'Unassigned'
                - 'due_date': Only if explicitly mentioned with a specific date, otherwise 'Not specified'
                - 'context': Short quote from transcript showing where task was mentioned
                
                Meeting Transcript:
                {transcript}
                
                Team Members (ONLY these people can be assigned tasks):
                {participant_info}
                
                If someone is mentioned in the transcript but isn't in this team list, DO NOT assign tasks to them."""}
            ],
            max_tokens=1000,
            response_format={"type": "json_object"}
        )
        
        result = json.loads(response.choices[0].message.content)
        
        # Ensure we have a "tasks" property that is a list
        if "tasks" not in result or not isinstance(result["tasks"], list):
            result = {"tasks": []}
        
        # Validation step: only keep tasks assigned to actual participants
        validated_tasks = []
        for task in result["tasks"]:
            # Convert assignee to lowercase for comparison
            assignee_lower = task.get("assignee", "").lower()
            
            # Check if:
            # 1. The task has actual content
            # 2. Either the assignee is "Unassigned" or matches someone in our participant list
            if task.get("task", "").strip() and (
                assignee_lower == "unassigned" or assignee_lower in participant_names
            ):
                # Add email based on participant list
                if assignee_lower in name_to_email:
                    task["email"] = name_to_email[assignee_lower]
                else:
                    task["email"] = ""
                
                validated_tasks.append(task)
        
        return {"tasks": validated_tasks}
    
    except Exception as e:
        print(f"Error extracting tasks: {str(e)}")
        return {"tasks": []}

# Function to generate email with tasks
def generate_task_email(summary, tasks, person_name):
    html = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 650px;
                margin: 0 auto;
            }}
            .header {{
                background-color: #1E40AF;
                color: white;
                padding: 20px;
                text-align: center;
                border-radius: 5px 5px 0 0;
            }}
            .content {{
                padding: 20px;
                background-color: #f9f9f9;
                border: 1px solid #ddd;
            }}
            .summary {{
                background-color: #EFF6FF;
                padding: 15px;
                border-left: 5px solid #3B82F6;
                margin-bottom: 20px;
            }}
            .task {{
                background-color: #F3F4F6;
                padding: 15px;
                margin-bottom: 10px;
                border-left: 5px solid #6B7280;
            }}
            .footer {{
                text-align: center;
                padding: 15px;
                font-size: 0.8em;
                color: #666;
                border-top: 1px solid #ddd;
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h2>Meeting Summary & Action Items</h2>
        </div>
        <div class="content">
            <p>Hello {person_name},</p>
            
            <p>Below you'll find a summary of our recent meeting and your assigned action items.</p>
            
            <h3>Meeting Summary</h3>
            <div class="summary">
                {summary.replace('\n', '<br>')}
            </div>
            
            <h3>Your Action Items</h3>
            """
    
    # Add tasks to email
    if tasks:
        for task in tasks:
            due_date = task.get('due_date', 'Not specified')
            task_desc = task.get('task', 'No description')
            
            html += f"""
            <div class="task">
                <strong>Task:</strong> {task_desc}<br>
                <strong>Due Date:</strong> {due_date}
            </div>
            """
    else:
        html += "<p>No specific action items were assigned to you from this meeting.</p>"
    
    # Close HTML
    html += f"""
            <p>Best regards,<br>Meeting Coordinator</p>
        </div>
        <div class="footer">
            This email was automatically generated by the Meeting Minutes Analyzer.
        </div>
    </body>
    </html>
    """
    
    return html

# Keeping the original summary-only email function intact
def generate_summary_email(summary, person_name):
    html = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 650px;
                margin: 0 auto;
            }}
            .header {{
                background-color: #1E40AF;
                color: white;
                padding: 20px;
                text-align: center;
                border-radius: 5px 5px 0 0;
            }}
            .content {{
                padding: 20px;
                background-color: #f9f9f9;
                border: 1px solid #ddd;
            }}
            .summary {{
                background-color: #EFF6FF;
                padding: 15px;
                border-left: 5px solid #3B82F6;
                margin-bottom: 20px;
            }}
            .footer {{
                text-align: center;
                padding: 15px;
                font-size: 0.8em;
                color: #666;
                border-top: 1px solid #ddd;
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h2>Meeting Summary</h2>
        </div>
        <div class="content">
            <p>Hello {person_name},</p>
            
            <p>Below you'll find a summary of our recent meeting.</p>
            
            <h3>Meeting Summary</h3>
            <div class="summary">
                {summary.replace('\n', '<br>')}
            </div>
            
            <p>Best regards,<br>Meeting Coordinator</p>
        </div>
        <div class="footer">
            This email was automatically generated by the Meeting Minutes Analyzer.
        </div>
    </body>
    </html>
    """
    
    return html

# Check for auth code in URL - FIXED with st.query_params
auth_code_from_url = check_url_for_auth_code()
if auth_code_from_url and not st.session_state.authenticated:
    with st.spinner("Authenticating with Gmail..."):
        # Make sure flow is initialized if coming directly with code in URL
        if st.session_state.flow is None:
            start_auth_flow()
        
        creds = complete_auth_flow(auth_code_from_url)
        if creds:
            st.session_state.service = build('gmail', 'v1', credentials=creds)
            st.session_state.authenticated = True
            # Clear the URL parameters - FIXED
            st.query_params.clear()
            st.success("Authentication successful!")

# Check if service is already authenticated
if not st.session_state.authenticated:
    existing_service = get_gmail_service()
    if existing_service:
        st.session_state.service = existing_service
        st.session_state.authenticated = True

# App UI
st.title("🔍 Meeting Minutes Analyzer")
st.markdown("Upload meeting transcripts, analyze tasks, and automatically email participants with their action items.")

# Sidebar for authentication
with st.sidebar:
    st.header("Gmail Authentication")
    
    if not st.session_state.authenticated:
        st.info("Please authenticate with Gmail to enable email sending functionality")
        
        if st.session_state.auth_url:
            st.markdown(f"1. [Click here to authorize with Gmail]({st.session_state.auth_url})")
            st.markdown("2. After authorization, copy the authorization code if not redirected properly")
            auth_code = st.text_input("Enter the authorization code (if needed):")
            
            if st.button("Submit Authorization Code"):
                with st.spinner("Completing authentication..."):
                    creds = complete_auth_flow(auth_code)
                    if creds:
                        st.session_state.service = build('gmail', 'v1', credentials=creds)
                        st.session_state.authenticated = True
                        st.success("Authentication successful!")
                        st.rerun()
        else:
            if st.button("Login to Gmail"):
                with st.spinner("Preparing authentication..."):
                    auth_url = start_auth_flow()
                    st.rerun()
    else:
        st.markdown('<div class="success-box">✅ Connected to Gmail</div>', unsafe_allow_html=True)
        if st.button("Logout"):
            if os.path.exists(TOKEN_FILE):
                os.remove(TOKEN_FILE)
            st.session_state.authenticated = False
            st.session_state.service = None
            st.session_state.auth_url = None
            st.session_state.flow = None
            st.rerun()
    
    st.markdown("---")
    st.markdown("### About This App")
    st.markdown("""
    This application helps you:
    
    1. **Analyze meeting transcripts** to extract key points
    2. **Identify action items** from the discussion
    3. **Assign tasks** to the appropriate team members
    4. **Email participants** with their assigned responsibilities
    
    Upload your meeting transcript and participant list to get started.
    """)

# Main area
tab1, tab2, tab3 = st.tabs(["📤 Upload Files", "📋 Analysis & Tasks", "📬 Send Emails"])

# Upload Files Tab
with tab1:
    st.header("Upload Meeting Documents")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Meeting Transcript")
        transcript_file = st.file_uploader("Upload meeting transcript", type=["txt", "docx", "pdf"], key="transcript_uploader")
        
        if transcript_file is not None:
            transcript_content = read_transcript(transcript_file)
            if transcript_content:
                st.session_state.transcript_content = transcript_content
                st.success(f"Successfully read transcript: {transcript_file.name}")
                with st.expander("Preview Transcript"):
                    st.text_area("Transcript Content", transcript_content, height=300)
    
    with col2:
        st.subheader("Participants List")
        st.markdown("""
        Upload a file with participant information in the format:
        - CSV: columns named 'name', 'email', 'expertise'
        - TXT/DOCX: each line with 'name, email, expertise'
        """)
        
        participants_file = st.file_uploader("Upload participants list", type=["csv", "txt", "docx"], key="participants_uploader")
        
        if participants_file is not None:
            participants = read_participants(participants_file)
            if participants:
                st.session_state.participants = participants
                st.success(f"Successfully read {len(participants)} participants from {participants_file.name}")
                with st.expander("Preview Participants"):
                    for p in participants:
                        st.markdown(f"**{p['name']}** - {p['email']} - *{p['expertise']}*")
    
    # Sample files for demonstration
    st.markdown("---")
    with st.expander("Need sample files?"):
        st.markdown("""
        **Sample Participant List Format (CSV):**
        ```
        name,email,expertise,role
        John Smith,john@example.com,Marketing,Marketing Director
        Sarah Johnson,sarah@example.com,Development,Lead Developer
        Ahmed Khan,ahmed@example.com,Finance,Financial Analyst
        ```
        
        **Sample Transcript Format:**
        ```
        Meeting Title: Q2 Planning Session
        Date: May 4, 2023
        
        Alex (Meeting Chair): Good morning everyone, let's get started with our Q2 planning session.
        
        John: I've prepared the marketing forecast as requested. We need to finalize the campaign budget by next Friday.
        
        Sarah: The development team will need two weeks to implement the new features we discussed.
        
        Ahmed: I'll prepare the financial projections by Wednesday and share them with the team.
        
        Alex: Great. John, please coordinate with Sarah on the launch timeline. Ahmed, we'll need those projections before Monday's executive meeting.
        ```
        """)

# Analysis & Tasks Tab
with tab2:
    st.header("Meeting Analysis & Task Assignment")
    
    # Check if both transcript and participants are loaded
    if st.session_state.transcript_content and st.session_state.participants:
        if st.button("Analyze Meeting Transcript"):
            with st.spinner("Analyzing meeting transcript and extracting tasks..."):
                # Generate summary
                summary = generate_meeting_summary(st.session_state.transcript_content)
                st.session_state.summary = summary
                
                # Extract and assign tasks
                tasks_result = extract_tasks_and_assign(st.session_state.transcript_content, st.session_state.participants)
                st.session_state.tasks = tasks_result
                
                st.success("Analysis complete!")
        
        # Display results if available
        if st.session_state.summary:
            st.subheader("Meeting Summary")
            st.markdown(f'<div class="summary-box">{st.session_state.summary}</div>', unsafe_allow_html=True)
        
        if st.session_state.tasks and "tasks" in st.session_state.tasks:
            if not st.session_state.tasks["tasks"]:
                st.info("No explicit tasks or action items were identified in this meeting transcript.")
            else:
                st.subheader("Task Assignments")
                
                # Group tasks by assignee
                tasks_by_assignee = {}
                for task in st.session_state.tasks["tasks"]:
                    assignee = task.get("assignee", "Unassigned")
                    if assignee not in tasks_by_assignee:
                        tasks_by_assignee[assignee] = []
                    tasks_by_assignee[assignee].append(task)
                
                # Display tasks by assignee
                for assignee, tasks in tasks_by_assignee.items():
                    with st.expander(f"{assignee} - {len(tasks)} tasks"):
                        for task in tasks:
                            st.markdown(f"""
                            <div class="task-box">
                                <strong>Task:</strong> {task.get('task', 'No description')}<br>
                                <strong>Due Date:</strong> {task.get('due_date', 'Not specified')}
                            </div>
                            """, unsafe_allow_html=True)
    else:
        st.info("Please upload both a meeting transcript and participants list in the Upload Files tab.")

# Send Emails Tab
with tab3:
    st.header("Send Meeting Information to Participants")
    
    if not st.session_state.authenticated:
        st.warning("Please authenticate with Gmail in the sidebar to enable email sending.")
    
    elif st.session_state.summary:
        st.subheader("Preview and Send Emails")
        
        # Group tasks by assignee email if tasks exist
        tasks_by_email = {}
        email_to_name = {}
        
        if st.session_state.tasks and "tasks" in st.session_state.tasks and st.session_state.tasks["tasks"]:
            for task in st.session_state.tasks["tasks"]:
                email = task.get("email", "")
                assignee = task.get("assignee", "Unassigned")
                
                # Skip if email is empty or invalid
                if not email or "@" not in email:
                    continue
                    
                if email not in tasks_by_email:
                    tasks_by_email[email] = []
                tasks_by_email[email].append(task)
                email_to_name[email] = assignee
        
        # Check if we have any valid tasks with emails
        has_tasks = bool(tasks_by_email)
        
        # For summary-only emails (when no tasks were found)
        if not has_tasks and st.session_state.participants:
            st.info("No specific tasks were identified in the transcript. You can still send the meeting summary to all participants.")
            
            # Create a mapping of all participants' emails
            for participant in st.session_state.participants:
                email = participant.get("email", "")
                name = participant.get("name", "")
                
                if email and "@" in email:
                    email_to_name[email] = name
            
            # Preview summary-only email
            if email_to_name:
                example_email = next(iter(email_to_name.items()))
                email, name = example_email
                
                st.markdown("### Summary Email Preview")
                email_content = generate_task_email(st.session_state.summary, name)
                st.components.v1.html(email_content, height=500, scrolling=True)
                
                st.markdown(f"This summary will be sent to all {len(email_to_name)} participants.")
                
                # Send all summary emails button
                if st.button("Send Summary to All Participants"):
                    success_count = 0
                    error_count = 0
                    
                    with st.spinner("Sending emails to participants..."):
                        progress_bar = st.progress(0)
                        
                        for i, (email, name) in enumerate(email_to_name.items()):
                            email_content = generate_task_email(st.session_state.summary, name)
                            
                            success, result = send_email(
                                st.session_state.service, 
                                email, 
                                "Meeting Summary", 
                                email_content
                            )
                            
                            if success:
                                success_count += 1
                            else:
                                error_count += 1
                            
                            # Update progress bar
                            progress_bar.progress((i + 1) / len(email_to_name))
                        
                        if success_count > 0:
                            st.success(f"✅ Successfully sent {success_count} summary emails")
                        
                        if error_count > 0:
                            st.error(f"❌ Failed to send {error_count} emails")
            else:
                st.warning("No valid participant email addresses found. Please check your participants file.")
                
        # If we have tasks, show the original task-based email interface
        elif has_tasks:
            # Create email preview tabs
            email_tabs = st.tabs([f"{name} ({email})" for email, name in email_to_name.items()])
            
            for i, (email, tasks) in enumerate(tasks_by_email.items()):
                with email_tabs[i]:
                    name = email_to_name[email]
                    email_content = generate_summary_email(st.session_state.summary, tasks, name)
                    
                    st.markdown("### Email Preview")
                    st.components.v1.html(email_content, height=500, scrolling=True)
            
            # Send all emails button
            if st.button("Send All Task Emails"):
                success_count = 0
                error_count = 0
                
                with st.spinner("Sending emails to participants..."):
                    progress_bar = st.progress(0)
                    
                    for i, (email, tasks) in enumerate(tasks_by_email.items()):
                        name = email_to_name[email]
                        email_content = generate_summary_email(st.session_state.summary, tasks, name)
                        
                        success, result = send_email(
                            st.session_state.service, 
                            email, 
                            "Meeting Action Items", 
                            email_content
                        )
                        
                        if success:
                            success_count += 1
                        else:
                            error_count += 1
                        
                        # Update progress bar
                        progress_bar.progress((i + 1) / len(tasks_by_email))
                    
                    if success_count > 0:
                        st.success(f"✅ Successfully sent {success_count} emails")
                    
                    if error_count > 0:
                        st.error(f"❌ Failed to send {error_count} emails")
        else:
            st.warning("No tasks identified and no participant list provided. Please upload a participant list to send the summary.")
    else:
        st.info("Please analyze the meeting transcript in the Analysis & Tasks tab before sending emails.")
